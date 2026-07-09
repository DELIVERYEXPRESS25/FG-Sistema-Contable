"""Genera el Manual Contable del Sistema F&G en formato Word."""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Estilos base ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1A, 0x1D, 0x2E)

def add_table(doc, headers, rows, col_widths=None):
    """Agrega una tabla con headers y filas."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Headers
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

def add_code_block(doc, text):
    """Agrega un bloque de código con fondo gris."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1D, 0x2E)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pf = p.paragraph_format
    pf.left_indent = Cm(1)
    return p

# ══════════════════════════════════════════════════════════════
#  PORTADA
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('F & G — Sistema Contable')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x4F, 0x8C, 0xFF)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MANUAL CONTABLE')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x1D, 0x2E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Documentación completa del sistema contable,\nplan de cuentas, asientos automáticos, inventario PEPS\ny reportes financieros.')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x7A, 0x7F, 0x99)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Versión 1.0 — Junio 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x7A, 0x7F, 0x99)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  CONTENIDO
# ══════════════════════════════════════════════════════════════
doc.add_heading('Contenido', level=1)
toc_items = [
    '1. Introducción',
    '2. Plan de Cuentas Jerárquico',
    '3. Reglas de Saldo (Debe / Haber)',
    '4. Flujo de una Venta POS',
    '5. Método PEPS (Primero en Entrar, Primero en Salir)',
    '6. Asientos Contables Automáticos',
    '7. Cierre Mensual',
    '8. Reportes Contables',
    '9. Ejemplo Práctico Completo',
    '10. Fórmulas Clave',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  1. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════
doc.add_heading('1. Introducción', level=1)
doc.add_paragraph(
    'El Sistema Contable F&G es una aplicación web desarrollada para gestionar las operaciones '
    'contables de una empresa comercial. El sistema automatiza el registro de operaciones, '
    'el control de inventario por el método PEPS, la generación de asientos contables y '
    'la elaboración de estados financieros.'
)
doc.add_paragraph('Características principales:')
items = [
    'Plan de cuentas jerárquico con códigos como 1.1.01, 4.1.01, 5.1, etc.',
    'Asientos contables automáticos al registrar ventas (POS) y movimientos de kardex.',
    'Inventario por el método PEPS (Primero en Entrar, Primero en Salir).',
    'Reportes financieros: Balanza de Comprobación, Estado de Resultados, Balance General.',
    'Cierre mensual automático con traspaso de resultados a cuentas de capital.',
    'Auditoría completa de todas las operaciones.',
    'Exportación a Excel con hojas separadas por reporte.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  2. PLAN DE CUENTAS
# ══════════════════════════════════════════════════════════════
doc.add_heading('2. Plan de Cuentas Jerárquico', level=1)
doc.add_paragraph(
    'El plan de cuentas utiliza códigos jerárquicos con punto como separador. '
    'Cada nivel representa una agrupación más específica. Por ejemplo, el código 1.1.01 '
    'significa: Activo (1) → Activo Corriente (1.1) → Efectivo (1.1.01).'
)

doc.add_heading('2.1 Estructura Completa', level=2)
cuentas_data = [
    ('1', 'Activo', 'Activo', '0'),
    ('1.1', 'Activo Corriente', 'Activo', '1'),
    ('1.1.01', 'Efectivo', 'Activo', '2'),
    ('1.1.02', 'Bancos', 'Activo', '2'),
    ('1.1.03', 'Cuentas por cobrar', 'Activo', '2'),
    ('1.1.04', 'Inventario de mercancías', 'Activo', '2'),
    ('1.1.05', 'Deudores Diversos', 'Activo', '2'),
    ('1.2', 'Activo No Corriente', 'Activo', '1'),
    ('1.2.01', 'Terreno', 'Activo', '2'),
    ('1.2.02', 'Edificio', 'Activo', '2'),
    ('1.2.03', 'Mobiliario y Equipo', 'Activo', '2'),
    ('1.2.04', 'Equipo de Cómputo Electrónico', 'Activo', '2'),
    ('2', 'Pasivo', 'Pasivo', '0'),
    ('2.1', 'Pasivo Corriente', 'Pasivo', '1'),
    ('2.1.01', 'Proveedores', 'Pasivo', '2'),
    ('2.1.02', 'Acreedores Diversos', 'Pasivo', '2'),
    ('2.1.03', 'Impuestos por Pagar', 'Pasivo', '2'),
    ('2.2', 'Pasivo No Corriente', 'Pasivo', '1'),
    ('2.2.01', 'Préstamos Bancarios Por Pagar LP', 'Pasivo', '2'),
    ('3', 'Patrimonio', 'Capital', '0'),
    ('3.1', 'Capital Social', 'Capital', '1'),
    ('3.2', 'Capital Contable', 'Capital', '1'),
    ('3.3', 'Resultados', 'Capital', '1'),
    ('3.3.01', 'Utilidad del Ejercicio', 'Capital', '2'),
    ('3.3.02', 'Pérdida del Ejercicio', 'Capital', '2'),
    ('3.4', 'Utilidad Acumulada', 'Capital', '1'),
    ('4', 'Ingresos', 'Ingreso', '0'),
    ('4.1', 'Ingresos por Ventas', 'Ingreso', '1'),
    ('4.1.01', 'Ventas al contado', 'Ingreso', '2'),
    ('4.1.02', 'Ventas al crédito', 'Ingreso', '2'),
    ('4.2', 'Devoluciones sobre Ventas', 'Ingreso', '1'),
    ('5', 'Costos y Gastos', 'Gasto', '0'),
    ('5.1', 'Costo de Ventas', 'Gasto', '1'),
    ('5.2', 'Gastos de Operación', 'Gasto', '1'),
    ('5.2.01', 'Sueldos y Salarios', 'Gasto', '2'),
    ('5.2.02', 'Renta del Local', 'Gasto', '2'),
    ('5.2.03', 'Servicios Básicos (Luz, agua, internet)', 'Gasto', '2'),
    ('5.2.04', 'Publicidad y Marketing', 'Gasto', '2'),
    ('5.2.05', 'Papelería y Empaques', 'Gasto', '2'),
]
add_table(doc, ['Código', 'Nombre', 'Tipo', 'Nivel'], cuentas_data)
doc.add_paragraph()

doc.add_heading('2.2 Niveles del Plan de Cuentas', level=2)
doc.add_paragraph(
    'Nivel 0: Cuentas de control (1=Activo, 2=Pasivo, 3=Patrimonio, 4=Ingresos, 5=Gastos).\n'
    'Nivel 1: Subgrupos (1.1=Activo Corriente, 4.1=Ingresos por Ventas, etc.).\n'
    'Nivel 2: Cuentas detalladas donde se registran los movimientos (1.1.01=Efectivo, 4.1.01=Ventas al contado).'
)
doc.add_paragraph(
    'Las cuentas de nivel 0 y 1 son "cuentas padre" que agrupan subcuentas. '
    'Solo las cuentas de nivel 2 (hojas) reciben movimientos directos.'
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  3. REGLAS DE SALDO
# ══════════════════════════════════════════════════════════════
doc.add_heading('3. Reglas de Saldo (Debe / Haber)', level=1)
doc.add_paragraph(
    'Cada tipo de cuenta tiene una regla para determinar si su saldo es deudor o acreedor:'
)
add_table(doc, [
    'Tipo de Cuenta',
    'Regla de Saldo',
    'Saldo Normal',
    'Ejemplo',
], [
    ['Activo', 'Saldo = Debe − Haber', 'Deudor', 'Efectivo: si recibe C$100 y paga C$30, saldo = C$70 Deudor'],
    ['Gasto', 'Saldo = Debe − Haber', 'Deudor', 'Renta: si se registra C$5,000 en Debe, saldo = C$5,000 Deudor'],
    ['Pasivo', 'Saldo = Haber − Debe', 'Acreedor', 'Proveedores: si se debe C$2,000, saldo = C$2,000 Acreedor'],
    ['Capital', 'Saldo = Haber − Debe', 'Acreedor', 'Capital Social: aportes van al Haber'],
    ['Ingreso', 'Saldo = Haber − Debe', 'Acreedor', 'Ventas: C$10,000 en Haber, saldo = C$10,000 Acreedor'],
])
doc.add_paragraph()

doc.add_heading('3.1 Fórmula en el Sistema', level=2)
add_code_block(doc, 'def tipo_saldo(tipo_cuenta):\n'
                    '    return "Debe" if tipo_cuenta in ["Activo", "Gasto"] else "Haber"')
doc.add_paragraph(
    'Cuando tipo_saldo es "Debe": saldo_debe = max(debe - haber, 0), saldo_haber = max(haber - debe, 0).\n'
    'Cuando tipo_saldo es "Haber": saldo_haber = max(haber - debe, 0), saldo_debe = max(debe - haber, 0).'
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  4. FLUJO DE UNA VENTA POS
# ══════════════════════════════════════════════════════════════
doc.add_heading('4. Flujo de una Venta POS', level=1)
doc.add_paragraph(
    'Cuando se realiza una venta en el Punto de Venta (POS), el sistema ejecuta '
    'automáticamente los siguientes pasos:'
)

steps = [
    ('Paso 1: Validar período', 'Verifica que la fecha de la venta no esté en un período cerrado.'),
    ('Paso 2: Procesar cada línea de venta', 'Para cada producto:\n'
     '  • Calcula subtotal = cantidad × precio de venta\n'
     '  • Obtiene el costo unitario desde el inventario PEPS\n'
     '  • Ejecuta la salida PEPS (consume lotes en orden FIFO)\n'
     '  • Registra el movimiento de kardex (salida)'),
    ('Paso 3: Calcular totales', 'total_venta = suma de subtotales\ntotal_costo = suma de costos de los productos vendidos'),
    ('Paso 4: Crear asiento diario automático', 'Genera un asiento con 2 o 4 movimientos (ver sección 6)'),
    ('Paso 5: Registrar en caja', 'Si el pago es en efectivo o banco, registra el movimiento de caja'),
    ('Paso 6: Registrar cuentas por cobrar', 'Si el pago es a crédito, crea un registro pendiente de cobro'),
    ('Paso 7: Guardar en historial', 'Almacena la venta completa en pos_historial con ref, fecha, líneas, totales'),
]
for title, desc in steps:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  5. MÉTODO PEPS
# ══════════════════════════════════════════════════════════════
doc.add_heading('5. Método PEPS (Primero en Entrar, Primero en Salir)', level=1)
doc.add_paragraph(
    'El método PEPS (o FIFO en inglés) es un sistema de costeo de inventario donde '
    'los primeros productos en entrar son los primeros en salir. Esto significa que '
    'el costo de los productos vendidos se basa en el costo de los lotes más antiguos.'
)

doc.add_heading('5.1 Estructura de Lotes', level=2)
doc.add_paragraph('Cada producto tiene una lista de lotes:')
add_code_block(doc, '{\n'
                    '    "lotes": [\n'
                    '        {"fecha": "2025-10-01", "cantidad": 50, "costo_unitario": 100.00, "cantidad_restante": 35},\n'
                    '        {"fecha": "2025-10-15", "cantidad": 30, "costo_unitario": 120.00, "cantidad_restante": 30}\n'
                    '    ],\n'
                    '    "stock_total": 65,\n'
                    '    "costo_promedio": 109.23\n'
                    '}')

doc.add_heading('5.2 Ejemplo de Salida PEPS', level=2)
doc.add_paragraph(
    'Si se venden 40 unidades:\n\n'
    '  Lote 1 (más antiguo): 35 unidades × C$100 = C$3,500 → Lote agotado\n'
    '  Lote 2 (siguiente):    5 unidades × C$120 = C$600 → Quedan 25 unidades\n\n'
    '  Costo total de la venta = C$3,500 + C$600 = C$4,100\n'
    '  Costo unitario promedio = C$4,100 ÷ 40 = C$102.50'
)

doc.add_heading('5.3 Costo Promedio Ponderado', level=2)
doc.add_paragraph(
    'Después de cada salida, se recalcula el costo promedio:\n\n'
    '  costo_promedio = Σ(cantidad_restante_i × costo_unitario_i) ÷ Σ(cantidad_restante_i)\n\n'
    'Este promedio solo incluye las unidades que quedan en stock, no las ya vendidas.'
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  6. ASIENTOS CONTABLES AUTOMÁTICOS
# ══════════════════════════════════════════════════════════════
doc.add_heading('6. Asientos Contables Automáticos', level=1)
doc.add_paragraph(
    'El sistema genera asientos contables automáticamente en las siguientes situaciones:'
)

doc.add_heading('6.1 Venta al Contado (con inventario)', level=2)
doc.add_paragraph('Asiento con 4 movimientos:')
add_table(doc, ['Cuenta', 'Tipo', 'Monto', 'Descripción'], [
    ['1.1.01 Efectivo', 'Debe', 'Total Venta', 'Entra dinero por la venta'],
    ['4.1.01 Ventas al contado', 'Haber', 'Total Venta', 'Ingreso por ventas'],
    ['5.1 Costo de Ventas', 'Debe', 'Total Costo', 'Costo de lo vendido'],
    ['1.1.04 Inventario de mercancías', 'Haber', 'Total Costo', 'Sale inventario'],
])
doc.add_paragraph()

doc.add_heading('6.2 Venta a Crédito (con inventario)', level=2)
doc.add_paragraph('Asiento con 4 movimientos:')
add_table(doc, ['Cuenta', 'Tipo', 'Monto', 'Descripción'], [
    ['1.1.03 Cuentas por cobrar', 'Debe', 'Total Venta', 'Cliente debe pagar'],
    ['4.1.01 Ventas al contado', 'Haber', 'Total Venta', 'Ingreso por ventas'],
    ['5.1 Costo de Ventas', 'Debe', 'Total Costo', 'Costo de lo vendido'],
    ['1.1.04 Inventario de mercancías', 'Haber', 'Total Costo', 'Sale inventario'],
])
doc.add_paragraph()

doc.add_heading('6.3 Venta por Banco (con inventario)', level=2)
doc.add_paragraph('Mismo estructura que contado, pero con cuenta 1.1.02 Bancos en lugar de 1.1.01.')
doc.add_paragraph()

doc.add_heading('6.4 Anulación de Venta', level=2)
doc.add_paragraph('El sistema genera un asiento de reversión (ref: ANUL-POS-XXXX):')
add_table(doc, ['Cuenta', 'Tipo', 'Monto', 'Descripción'], [
    ['4.1.01 Ventas al contado', 'Debe', 'Total Venta', 'Reversa el ingreso'],
    ['1.1.01/1.1.02/1.1.03', 'Haber', 'Total Venta', 'Reversa el cobro'],
    ['1.1.04 Inventario de mercancías', 'Debe', 'Total Costo', 'Devuelve inventario'],
    ['5.1 Costo de Ventas', 'Haber', 'Total Costo', 'Reversa el costo'],
])
doc.add_paragraph(
    'Además, el sistema revierte los movimientos de kardex, devuelve el stock '
    'a los lotes PEPS y elimina los registros de caja y cuentas por cobrar.'
)
doc.add_paragraph()

doc.add_heading('6.5 Referencias Automáticas', level=2)
add_table(doc, ['Referencia', 'Origen', 'Ejemplo'], [
    ['POS-XXXX', 'Venta POS', 'POS-0001, POS-0002'],
    ['KX-XXXX', 'Movimiento de kardex', 'KX-0001'],
    ['ANUL-POS-XXXX', 'Anulación de venta', 'ANUL-POS-0001'],
    ['CIERRE-TIPO-PERIODO', 'Cierre mensual', 'CIERRE-MENS-2026-07'],
    ['AJ', 'Ajuste contable', 'AJ'],
])
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  7. CIERRE MENSUAL
# ══════════════════════════════════════════════════════════════
doc.add_heading('7. Cierre Mensual', level=1)
doc.add_paragraph(
    'El cierre mensual es el proceso contable que transfiere los saldos de las cuentas de '
    'resultado (Ingresos y Gastos) a las cuentas de capital (Utilidad o Pérdida del Ejercicio). '
    'Esto "cierra" las cuentas de resultado dejándolas en cero para el siguiente período.'
)

doc.add_heading('7.1 Proceso de Cierre', level=2)
steps_cierre = [
    'El sistema identifica todas las cuentas de Ingreso y Gasto con movimientos en el período.',
    'Para cada cuenta de Ingreso con saldo acreedor (haber > debe): genera un Debe por el saldo para anularla.',
    'Para cada cuenta de Gasto con saldo deudor (debe > haber): genera un Haber por el saldo para anularla.',
    'Calcula la diferencia entre totales de Debe y Haber del asiento de cierre.',
    'Si hay utilidad (haber > debe): la registra en cuenta 3.3.01 Utilidad del Ejercicio (Haber).',
    'Si hay pérdida (debe > haber): la registra en cuenta 3.3.02 Pérdida del Ejercicio (Debe).',
    'Crea el asiento contable con ref: CIERRE-MENS-YYYY-MM.',
    'Marca el período como cerrado para evitar modificaciones.',
]
for i, step in enumerate(steps_cierre, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('7.2 Ejemplo de Cierre', level=2)
doc.add_paragraph(
    'Período: Octubre 2025\n'
    'Total Ingresos (4.1.01 + 4.1.02): C$183,185.00\n'
    'Total Gastos (5.1 + 5.2.01 + 5.2.02 + 5.2.03): C$59,800.00\n\n'
    'Utilidad = C$183,185 − C$59,800 = C$123,385.00'
)
doc.add_paragraph('Asiento de cierre:')
add_table(doc, ['Cuenta', 'Tipo', 'Monto'], [
    ['4.1.01 Ventas al contado', 'Debe', '183,185.00'],
    ['5.1 Costo de Ventas', 'Haber', '59,800.00'],
    ['5.2.01 Sueldos y Salarios', 'Haber', 'X,XXX.00'],
    ['5.2.02 Renta del Local', 'Haber', 'X,XXX.00'],
    ['3.3.01 Utilidad del Ejercicio', 'Haber', '123,385.00'],
])
doc.add_paragraph()

doc.add_heading('7.3 Tipos de Cierre Disponibles', level=2)
add_table(doc, ['Tipo', 'Frecuencia', 'Rango de Fechas'], [
    ['Mensual', 'Cada mes', 'Primer día al último día del mes'],
    ['Quincenal', 'Cada 15 días', 'Día 1-15 o día 16-fin de mes'],
    ['Semanal', 'Cada semana', 'Según semana ISO'],
])
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  8. REPORTES CONTABLES
# ══════════════════════════════════════════════════════════════
doc.add_heading('8. Reportes Contables', level=1)

doc.add_heading('8.1 Balanza de Comprobación', level=2)
doc.add_paragraph(
    'Verifica que los débitos equalen los créditos. Muestra todas las cuentas con movimientos, '
    'sus saldos deudores y acreedores.'
)
doc.add_paragraph('Columnas del Excel:')
add_table(doc, ['Columna', 'Descripción'], [
    ['Código', 'Código jerárquico de la cuenta'],
    ['Cuenta', 'Nombre de la cuenta'],
    ['Tipo', 'Activo, Pasivo, Capital, Ingreso o Gasto'],
    ['Saldo Inicial Deudor', 'Saldo deudor al inicio del período'],
    ['Saldo Inicial Acreedor', 'Saldo acreedor al inicio del período'],
    ['Mov. Deudor', 'Movimientos al Debe durante el período'],
    ['Mov. Acreedor', 'Movimientos al Haber durante el período'],
    ['Saldo Final Deudor', 'Saldo deudor al cierre del período'],
    ['Saldo Final Acreedor', 'Saldo acreedor al cierre del período'],
])
doc.add_paragraph()

doc.add_heading('8.2 Estado de Resultados', level=2)
doc.add_paragraph(
    'Muestra los ingresos, gastos y la utilidad o pérdida neta del período.\n\n'
    'Fórmula: Utilidad Neta = Total Ingresos − Total Gastos'
)
doc.add_paragraph('Estructura:')
doc.add_paragraph('  INGRESOS\n    Ventas al contado: C$XXX\n    Ventas al crédito: C$XXX\n    Total Ingresos: C$XXX\n\n'
                   '  GASTOS\n    Costo de Ventas: C$XXX\n    Sueldos: C$XXX\n    Renta: C$XXX\n    Total Gastos: C$XXX\n\n'
                   '  UTILIDAD NETA: C$XXX')
doc.add_paragraph()

doc.add_heading('8.3 Balance General', level=2)
doc.add_paragraph(
    'Muestra la ecuación contable fundamental:\n\n'
    '  ACTIVOS = PASIVOS + PATRIMONIO (Capital)\n\n'
    'Se divide en 3 secciones:\n'
    '  • Activos: todo lo que la empresa posee (efectivo, inventario, equipo)\n'
    '  • Pasivos: todo lo que la empresa debe (proveedores, préstamos)\n'
    '  • Capital: aportes de losDueños + Utilidad del Ejercicio'
)
doc.add_paragraph(
    'La utilidad del período se refleja automáticamente en las cuentas:\n'
    '  • 3.3.01 Utilidad del Ejercicio (si hay ganancia)\n'
    '  • 3.3.02 Pérdida del Ejercicio (si hay pérdida)'
)
doc.add_paragraph()

doc.add_heading('8.4 Libro Diario', level=2)
doc.add_paragraph(
    'Registro cronológico de todos los asientos contables. Muestra:\n'
    '  • Número de asiento, fecha, descripción, referencia\n'
    '  • Cuenta, nombre de cuenta, monto al Debe, monto al Haber\n'
    '  • Incluye asientos de ventas POS, ajustes y cierres mensuales'
)
doc.add_paragraph()

doc.add_heading('8.5 Libro Mayor', level=2)
doc.add_paragraph(
    'Resumen por cuenta de todos los movimientos. Muestra:\n'
    '  • Cuenta, nombre, tipo\n'
    '  • Saldo Inicial Deudor/Acreedor\n'
    '  • Debe del período, Haber del período\n'
    '  • Saldo Final Deudor/Acreedor'
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  9. EJEMPLO PRÁCTICO COMPLETO
# ══════════════════════════════════════════════════════════════
doc.add_heading('9. Ejemplo Práctico Completo', level=1)
doc.add_paragraph(
    'Ejemplo: La empresa F&G compra 100 anillos de acero a C$45 c/u y vende 20 al contado a C$65 c/u.'
)

doc.add_heading('9.1 Paso 1: Entrada de Inventario', level=2)
doc.add_paragraph(
    'Se registran 100 unidades a C$45 c/u.\n'
    'Costo total = 100 × C$45 = C$4,500'
)
doc.add_paragraph('Asiento:')
add_table(doc, ['Cuenta', 'Tipo', 'Monto'], [
    ['1.1.04 Inventario de mercancías', 'Debe', '4,500.00'],
    ['2.1.01 Proveedores', 'Haber', '4,500.00'],
])
doc.add_paragraph()

doc.add_heading('9.2 Paso 2: Venta al Contado', level=2)
doc.add_paragraph(
    'Se venden 20 unidades a C$65 c/u.\n'
    'Total venta = 20 × C$65 = C$1,300\n'
    'Total costo = 20 × C$45 = C$900 (PEPS: primer lote)'
)
doc.add_paragraph('Asiento:')
add_table(doc, ['Cuenta', 'Tipo', 'Monto', 'Descripción'], [
    ['1.1.01 Efectivo', 'Debe', '1,300.00', 'Cobro por venta'],
    ['4.1.01 Ventas al contado', 'Haber', '1,300.00', 'Ingreso por venta'],
    ['5.1 Costo de Ventas', 'Debe', '900.00', 'Costo de lo vendido'],
    ['1.1.04 Inventario de mercancías', 'Haber', '900.00', 'Sale inventario'],
])
doc.add_paragraph()

doc.add_heading('9.3 Estado del Inventario después de la venta', level=2)
doc.add_paragraph(
    'Lote 1: 80 unidades restantes × C$45 = C$3,600\n'
    'Stock total: 80 unidades\n'
    'Costo promedio: C$45.00'
)

doc.add_heading('9.4 Cálculo de Utilidad', level=2)
doc.add_paragraph(
    'Ingresos: C$1,300\n'
    'Gastos (Costo de Ventas): C$900\n'
    'Utilidad Neta: C$1,300 − C$900 = C$400\n\n'
    'Margen de ganancia: (C$400 ÷ C$1,300) × 100 = 30.77%'
)

doc.add_heading('9.5 Precio de Venta (fórmula)', level=2)
doc.add_paragraph(
    'Fórmula del sistema:\n'
    '  Precio Venta = Costo ÷ (1 − Margen%\n\n'
    'Ejemplo con margen del 30%:\n'
    '  Precio Venta = C$45 ÷ (1 − 0.30)\n'
    '  Precio Venta = C$45 ÷ 0.70\n'
    '  Precio Venta = C$64.29 ≈ C$65.00'
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  10. FÓRMULAS CLAVE
# ══════════════════════════════════════════════════════════════
doc.add_heading('10. Fórmulas Clave', level=1)

formulas = [
    ('Precio de Venta', 'PV = Costo ÷ (1 − Margen% / 100)', 'Determina el precio de venta deseando un margen de ganancia específico.'),
    ('Utilidad Neta', 'Utilidad = Total Ingresos − Total Gastos', 'Resultado del Estado de Resultados.'),
    ('Balance General', 'Activos = Pasivos + Capital', 'Ecuación contable fundamental. Debe cuadrar siempre.'),
    ('Costo Promedio PEPS', 'CP = Σ(Qi × Ci) ÷ Σ(Qi)', 'Promedio ponderado de los lotes disponibles en stock.'),
    ('Saldo Deudor', 'Saldo = Debe − Haber (si Debe > Habero)', 'Aplica para cuentas de Activo y Gasto.'),
    ('Saldo Acreedor', 'Saldo = Haber − Debe (si Haber > Debe)', 'Aplica para cuentas de Pasivo, Capital e Ingreso.'),
    ('Cierre de Ingresos', 'Se genera un Debe por el saldo del Haber', 'Anula la cuenta de Ingreso para traspasar a Utilidad.'),
    ('Cierre de Gastos', 'Se genera un Haber por el saldo del Debe', 'Anula la cuenta de Gasto para traspasar a Utilidad.'),
]

add_table(doc, ['Fórmula', 'Expresión', 'Descripción'], formulas)
doc.add_paragraph()

doc.add_heading('10.1 Notas Importantes', level=2)
doc.add_paragraph(
    '• La fórmula de precio de venta usa división (costo ÷ (1−margen)), NO multiplicación.\n'
    '  Esto asegura que el margen sea sobre el precio de venta, no sobre el costo.\n\n'
    '• Las cuentas de tipo "Activo" y "Gasto" tienen saldo deudor normal.\n'
    '  Las cuentas de tipo "Pasivo", "Capital" e "Ingreso" tienen saldo acreedor normal.\n\n'
    '• El asiento de cierre debe estar siempre balanceado: total Debe = total Haber.\n'
    '  La diferencia se registra en 3.3.01 (Utilidad) o 3.3.02 (Pérdida).\n\n'
    '• Una vez cerrado un período, no se pueden crear, editar ni eliminar asientos en ese período.'
)

# ── Guardar ──
output_path = os.path.join(os.path.dirname(__file__), 'Manual_Contable_FG.docx')
doc.save(output_path)
print(f'Documento guardado en: {output_path}')
